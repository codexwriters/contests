#!/usr/bin/env python
import docx
import os
import re
import zipfile
import random
import subprocess
import sendEmail as send
import uuid
from itertools import cycle
from datetime import date
from difflib import SequenceMatcher
from pathlib import Path


#dir = "/Users/leigh/Tresors/Codex"
dir = "/home/codex/development"
today = date.today()


def main():
    convertFiles()
    contestList, contestName = readContest()
    contestNoSpaces = formatContestTitle(contestName)
    ballotList, emailList = readEmail(contestList)
    originalExts = originalFileExtensionList(contestList)
    docxList = convertToDocx(originalExts)
    removeEmail(docxList)
    packetList = createPacket(contestList, emailList, contestNoSpaces, contestName)
    sendPacket(packetList, contestName)

def convertFiles():
    for filename in os.listdir(dir):
        if filename.endswith(".doc") or filename.endswith(".docx") \
        or filename.endswith(".rtf") or filename.endswith(".pdf") \
        or filename.endswith(".odt"):
            print("Converting", filename)
            command = "java -jar tika-app-1.24.1.jar --text " + filename + " > " + filename + ".txt"
            subprocess.run(command, shell=True)
        else:
            continue

def readContest():
    term = input("Enter contest: ")
    #contest = "(?<=Contest: " + term + ").*$"
    contest = "Contest: " + term
    #print(contest)
    contestList = []
    for filename in os.listdir(dir):
        if filename.endswith(".txt"):
            with open(filename) as file:
                while True:
                    line = file.readline()
                    # if not line:
                    #     break
                    while not line.startswith("Contest:"):
                        line = file.readline()
                    sequence = SequenceMatcher(a=contest, b=line)
                    comparisonLine = "\nContest is \"" + term + "\" and the contest" \
                    " found from story \"" + filename + "\" is " + line
                    #print(comparisonLine, end="")
                    if(sequence.ratio() > 0.75):
                        #print("Match percent is: ", sequence.ratio())
                        contestList.append(filename)
                        break
                    #if not re.search(contest, line):  # This could be handled more gracefully
                    else:
                        #print("Match percent is: ", sequence.ratio(), end="")
                        print("No match in", filename, "skipping file.\n")
                        break
    if len(contestList) == 0:
        print("No stories matching contest title found. Exiting.")
        quit()
    return(contestList, term)

def readEmail(contestList):
    emailList = []
    email = "(?<=Email: ).*$"
    for each in contestList:
        with open(each) as file:
            while True:
                line = file.readline()
                if not line:
                    break
                z = re.search(email, line)
                if z:
                    emailList.append(z.group(0))
                else:
                    continue
    random.shuffle(contestList)
    return(contestList, emailList)

def readTitle(fileName):
    title = "(?<=Title: ).*$"
    with open(fileName) as file:
        while True:
            line = file.readline()
            if not line:
                break
            z = re.search(title, line)
            if z:
                return z.group(0)
            else:
                continue

def originalFileExtensionList(contestList):
    originalExts = []
    for each in contestList:
        newEach = os.path.splitext(each)[0]
        originalExts.append(newEach)
    return originalExts

def convertToDocx(originalExts):
    docxList = []
    for each in originalExts:
        if each.endswith(".docx"):
            docxList.append(each)
        if each.endswith(".doc") or each.endswith(".rtf"):
            command = "lowriter --convert-to docx " + each
            subprocess.run(command, shell=True)
            nakedEach = os.path.splitext(each)[0]
            newDocx = nakedEach + ".docx"
            docxList.append(newDocx)
    return docxList

def removeEmail(docxList):
    for each in docxList:
        doc = docx.Document(each)
        try:
            all = doc.paragraphs
            header = all[0].text
            name = header.splitlines()[0]
            contest = header.splitlines()[1]
            title = header.splitlines()[3]
            fixedHeader = name + "\n" + contest + "\n" + title 
            all[0].text = fixedHeader
            doc.save(each)

        except:
            all = doc.paragraphs
            if doc.paragraphs[0].text.startswith("Email"):
                doc.paragraphs[0].text = ""
            if doc.paragraphs[1].text.startswith("Email"):
                doc.paragraphs[1].text = ""
            if doc.paragraphs[2].text.startswith("Email"):
                doc.paragraphs[2].text = ""
            if doc.paragraphs[3].text.startswith("Email"):
                doc.paragraphs[3].text = ""
            doc.save(each)


def formatContestTitle(contestName):
    noSpaces = re.compile(r'\s+')
    contestNoSpaces = re.sub(noSpaces, '', contestName)
    return contestNoSpaces

def createPacket(contestList, emailList, contestNoSpaces, contestName):
    for each in contestList:
        print(each)
    zippedList = []
    packetList = []
    attachments = []
    date = today.strftime("%m-%d-%Y")
    contestId = uuid.uuid4().hex
    packetSize = int(input("Enter packet size: "))
    #group = [docxList[i::packetSize] for i in range(packetSize)]
    group = [contestList[i::packetSize] for i in range(packetSize)]
    #group = [originalFiles[i::packetSize] for i in range(packetSize)]
    count = 0
    for every in group:
        count += 1
        fileName = contestNoSpaces + "_Packet_%s" % count + "_" + date
        ballotFileName = fileName + "_ballot.txt"
        with open(ballotFileName, "a") as f:
            f.write("Copy/paste the below into the word processor of your " + \
            "choice and, when finished, send to contest-bot as an attachment " + \
            "in a PM. Ratings are scored from 1 (worst) to 10 (best), according " + \
            "to how interesting or compelling the story was for you. Comments " + \
            "are optional. Please preserve the format of the file structure " + \
            "below, as the program needs to understand how to parse the " + \
            "results (including the file header). You may, however, save the " + \
            "document in any format--.pdf, .doc(x), .rtf, .txt, .odt are all " + \
            "acceptable.\n\n")
            f.write("File header:\n")
            f.write("Contest: " + contestName + "\n")
            f.write("ContestID: " + contestId + "\n")
            f.write("Date: " + date + "\n\n")
            f.write("VOTING STARTS HERE\n\n")
        #contestZipped = zipfile.ZipFile("/Users/leigh/Tresors/Codex/%s.zip" % fileName, mode='w', compression = zipfile.ZIP_DEFLATED)
        contestZipped = zipfile.ZipFile("/home/codex/development/%s.zip" % fileName, mode='w', compression = zipfile.ZIP_DEFLATED)
        for member in every:
            #newMember = os.path.splitext(member)[0]
            #newMember = member + ".txt"
            p = Path(member)
            extensions = "".join(p.suffixes)
            noExts = str(p).replace(extensions, "")
            docxExts = noExts + ".docx"
            title = readTitle(member)
            with open(ballotFileName, "a") as f:
                f.write("Title: " + title + " (" + member + ")" + "\n")
                f.write("Your rating: " + "\n")
                f.write("Your comments: " + "\n")
                f.write("-------------------------------------------------\n\n")
            contestZipped.write(docxExts)
        contestZipped.write(ballotFileName)
        contestZipped.close()
        zippedList.append(fileName)
    for i, a in enumerate(emailList):
        packetList.append((a, zippedList[i % len(zippedList)]))
    attachments.append(packetList[0][1])
    return(packetList)

def sendPacket(packetList, contestName):
    for each in packetList:
        attachments = []
        file = each[1] + ".zip"
        subject = "Codex Writers Gamma Contest: " + contestName
        print("Email", each[0], "will receive", file)
        attachments.append(file)
        body = "Hey there! <br>You're receiving this email because you submitted " \
        "a story for the current running contest. <br>" \
        "The attached file contains the stories that have been assigned to you, " \
        "as well as a ballot file. <br>" \
        "Please fill out the ballot file and send it as an attachment to @contestbot "\
        " on https://forums.codexgamma.com. <br>" \
        "At the appropriate time, we will process the ballots and send them to the authors." \
        "<br><br>" \
        "Please remember the following: <br>" \
        "- We're expecting bugs, but will be making fixes and improvements over time <br>" \
        "- This email address is unmonitored. If you have a problem, please ping @Leigh on the forums." \
        "<br><br>" \
        "Have fun! <br>" \
        "<br> Leigh"
        print(send.Send(subject, "contest@codexgamma.com", "Codex Writers Contest Bot", each[0], \
        body, "Text Body", True, attachments))
        #print(send.Send(subject, "contest@codexgamma.com", "Codex Writers Contest Bot", "leigh@regula.one", \
        #body, "Text Body", True, attachments))

main()
